import os
import io
import json

from pytz import timezone, utc
import boto3
from flask import make_response, jsonify
from bson import json_util
from .constants import BTS_EMAIL, BTS_APP_CONFIG_EMAIL, SGT_TIMEZONE
from exponent_server_sdk import (
    DeviceNotRegisteredError,
    PushClient,
    PushMessage,
    PushTicketError,
    PushServerError,
)

import rollbar

from requests.exceptions import ConnectionError, HTTPError

import functools
from flask_mail import Mail, Message
import tempfile
import xlsxwriter

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image


def parse_json(data):
    return json.loads(json_util.dumps(data))

def printJ(data):
    print(json.dumps(data, indent=4, sort_keys=False))


def upload_image(file_obj, bucket, file_name):
    """
    Function to upload a file to an S3 bucket
    """
    s3_client = boto3.client('s3')
    response = s3_client.upload_fileobj(file_obj, bucket, file_name)
    return response

def download_image(file_name, bucket):
    """
    Function to download a given file from an S3 bucket
    """
    s3 = boto3.client('s3')
    file_stream = io.BytesIO()
    s3.download_fileobj(bucket, file_name, file_stream)

    return file_stream

def list_images(bucket):
    """
    Function to list files in a given S3 bucket
    """
    s3 = boto3.client('s3')
    contents = []
    for item in s3.list_objects(Bucket=bucket)['Contents']:
        contents.append(item)

    return contents


def serverResponse(data, status_code, msg):
    packet = {
        "data": data,
        "description": msg
    }
    r = make_response(jsonify(parse_json(packet)))
    r.status_code = status_code
    return r

def send_push_message(token, title, message, extra=None):
    try:
        try:
            response = PushClient().publish(
                PushMessage(to=token,
                            body=message,
                            title=title,
                            data=extra,
                            sound="default"
                            )
            )
        except PushServerError as exc:
            # Encountered some likely formatting/validation error.
            rollbar.report_exc_info(
                extra_data={
                    'token': token,
                    'message': message,
                    'extra': extra,
                    'errors': exc.errors,
                    'response_data': exc.response_data,
                })
            raise
        except (ConnectionError, HTTPError) as exc:
            # Encountered some Connection or HTTP error - retry a few times in
            # case it is transient.
            rollbar.report_exc_info(
                extra_data={'token': token, 'message': message, 'extra': extra})
            raise self.retry(exc=exc)

        try:
            # We got a response back, but we don't know whether it's an error yet.
            # This call raises errors so we can handle them with normal exception
            # flows.
            response.validate_response()
        except DeviceNotRegisteredError:
            # Mark the push token as inactive
            # from notifications.models import PushToken
            # PushToken.objects.filter(token=token).update(active=False)
            print("DeviceNotRegisteredError")
        except PushTicketError as exc:
            # Encountered some other per-notification error.
            rollbar.report_exc_info(
                extra_data={
                    'token': token,
                    'message': message,
                    'extra': extra,
                    'push_response': exc.push_response._asdict(),
                })
            raise self.retry(exc=exc)
    except ValueError:
        pass

def utc2sgt(date):
    if date.tzinfo is None or date.tzinfo.utcoffset(date) is None:
        date = utc.localize(date)
    
    return date.astimezone(timezone(SGT_TIMEZONE))

# for checking data
def check_required_info(mydict, key_arr):
    missing_keys = []
    key_value_error = []

    if isinstance(mydict, dict) and isinstance(key_arr, list):
        for item in key_arr:
            # check if the dictionary contains keys
            if item not in mydict:
                missing_keys.append(item)

            else:
                if mydict[item] is None:
                    key_value_error.append(item)
                elif isinstance(mydict[item], str):
                    if len(mydict[item]) == 0:
                        key_value_error.append(item)

        return missing_keys, key_value_error
    else:
        return None

def validate_required_info(mydict, key_arr):
    if check_required_info(mydict, key_arr) is not None:
        missing_keys, key_value_error = check_required_info(mydict, key_arr)
        if len(missing_keys) == 0 and len(key_value_error) == 0:
            return True, ""

        else:
            error_message = {}

            if len(missing_keys) > 0:
                error_message["missing_keys"] = missing_keys

            if len(key_value_error) > 0:
                error_message["key_value_error"] = key_value_error

        return False, [error_message]

def check_duplicate(mongo, collection, key, value):
    results = mongo.db[collection].find({key: value}).count()
    if(results != 0):
        return True
    else:
        return False

def find_and_return_one(mongo, collection, key, value):
    try:
        result = mongo.db[collection].find_one({key: value})

        if result is not None:
            return True, result
        else:
            return False, None
    except:
        return None, None

# for excel operations
def nested_dict_get(dictionary, dotted_key):
    keys = dotted_key.split('.')
    return functools.reduce(lambda d, key: d.get(key) if d else None, keys, dictionary)

def excel_enter_field_down(worksheet, start_pos, field, val_dict):
    for i in range(len(field)):
        worksheet.write(start_pos[0]+i,
                        start_pos[1],
                        nested_dict_get(val_dict, field[i]))

def excel_enter_field_across(worksheet, start_pos, field, val_dict):
    for i in range(len(field)):
        worksheet.write(start_pos[0],
                        start_pos[1] + i,
                        nested_dict_get(val_dict, field[i]))

def from_audit_to_excel(workbook, page_name, data):
    # unpack data
    audit_dict = data["audit_info"]
    staff_dict = data["staff_info"]
    tenant_dict = data["tenant_info"]
    inst_dict = data["inst_info"]

    if("fnb" in audit_dict["auditChecklists"].keys()):
        form_type = "F&B"
    else:
        form_type = "Non-F&B"

    worksheet = workbook.add_worksheet(page_name)

    # Def param
    # pos = (row, col)
    staff_start_pos = (3, 1)
    inst_start_pos = (5, 1)
    tenant_start_pos = (3, 6)
    audit_start_pos = (11, 1)

    # Cell formats
    bold_underline = workbook.add_format({'bold': True, 'underline': True})
    bold = workbook.add_format({'bold': True})
    percentage = workbook.add_format()
    percentage.set_num_format('0.0%')

    # Def fields
    worksheet.write('A1', 'AUDIT INFORMATION', bold_underline)
    worksheet.set_column(0, 0, 18)
    worksheet.write('A3', 'Staff', bold)
    worksheet.write('A4', "Name :")
    worksheet.write('A5', "Email :")
    worksheet.write('A6', "Institution name:")
    worksheet.write('A7', "Institution acronym:")

    worksheet.write('F3', "Tenant", bold)
    worksheet.set_column(5, 5, 18)
    worksheet.write('F4', "Name :")
    worksheet.write('F5', "Email :")
    worksheet.write('F6', "Stall name :")

    worksheet.write('A9', "Date", bold)
    worksheet.write('A11', "Forms", bold)
    worksheet.write('B11', "Scores", bold)
    worksheet.write('C11', "Rectification Progress", bold)
    worksheet.set_column(1, 2, cell_format=percentage)

    # Value fields from dict
    staff_field = ["name", "email"]
    inst_field = ["name", "_id"]
    tenant_field = ["name", "email", "stallName"]
    audit_field = ["score", "rectificationProgress"]

    # get and enter field
    excel_enter_field_down(worksheet, staff_start_pos, staff_field, staff_dict)
    excel_enter_field_down(worksheet, inst_start_pos, inst_field, inst_dict)
    excel_enter_field_down(worksheet, tenant_start_pos,
                           tenant_field, tenant_dict)
    worksheet.write('B9', audit_dict["date"].strftime('%d/%m/%Y   %H:%M'))
    worksheet.write('A12', form_type)
    excel_enter_field_across(
        worksheet, audit_start_pos, audit_field, audit_dict)

    workbook.close()

    return workbook

# for word operations
def check_valid_param(data):
    if(len(data["missing"]) > 0) or (len(data["error"]) > 0):
        return False
    else:
        return True
    
def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width

def set_cell_background(cell, fill):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    from docx.oxml.shared import qn  # feel free to move these out
    from docx.oxml.xmlchemy import OxmlElement
    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath(
            'w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd')  # add new w:shd element to it
    if fill:
        # set fill property, respecting namespace
        cell_shading.set(qn('w:fill'), fill)
    # finally extend cell props with shading element
    cell_properties.append(cell_shading)

def insert_list_by_col(table, start_row, col_num, data_list):
    for i in range(len(data_list)):
        table.cell(start_row+i, col_num).text = data_list[i]

def insert_dict_by_col(table, start_row, col_num, data_list, data_dict):
    for i in range(len(data_list)):
        table.cell(start_row+i, col_num).text = data_dict[data_list[i]]

def make_rows_underline(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.underline = True

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
def table_center_align(*columns):
    for col in columns:
        for cell in col.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def get_rect(question_dict, ans_dict):
    rect_form = {}
    
    for checklist_name in question_dict.keys():        
        for i in range(len(question_dict[checklist_name])): 
            question = question_dict[checklist_name][i]['question']
            answer_bool = ans_dict[checklist_name][i]['answer']
            item = ans_dict[checklist_name][i]
            
            if answer_bool is None:
                pass
                
            elif answer_bool:
                pass
                
            else:
                rect_form.setdefault(checklist_name,{})
                
                #get images from db
                image_list = item.get("image", None)

                remarks = item.get("remarks", "-")
                rectified = item.get("rectified", False)

                deadline = item.get("deadline", None)
                if deadline is not None:
                    deadline =  deadline.strftime("%m/%d/%Y %H:%M:%S")
                else:
                    deadline = "-"
                if rectified:
                    rectified = "Yes"
                else:
                    rectified = "No"
                
                #get rectification images from db
                rect_remarks = item.get("rectificationRemarks", None)
                rect_image_list = item.get("rectificationImages", None)
                    
                #required information
                rect_form[checklist_name][question] = {
                    "Non-compliance Images": image_list,
                    "Remarks": remarks,
                    "Rectified": rectified,
                    "Deadline": deadline
                    }
                
                if rect_image_list is not None:
                    rect_form[checklist_name][question]["Rectification Images"] = rect_image_list
                
                if rect_remarks is not None:
                    rect_form[checklist_name][question]["Rectification Remarks"] = rect_remarks

    return rect_form
    
def map_qna(question_dict, ans_dict):
    form_with_ans = {}
    summary_table = []
    summary_form = {}
    total_points = 0
    
    for checklist_name in question_dict.keys():        
        form_with_ans[checklist_name] = []
        
        points = 0
        max_points = 0
        for i in range(len(question_dict[checklist_name])): 
            question = question_dict[checklist_name][i]['question']
            answer_bool = ans_dict[checklist_name][i]['answer']
            item = ans_dict[checklist_name][i]
            
            if answer_bool is None:
                pass
                
            elif answer_bool:
                form_with_ans[checklist_name].append({
                    "question": question,
                    "answer": "1"})
                points += 1
                max_points += 1
                
            else:
                if item["rectified"]:
                    form_with_ans[checklist_name].append({
                        "question": question,
                        "answer": "0*"})
                    max_points += 1
                else:
                    form_with_ans[checklist_name].append({
                        "question": question,
                        "answer": "0"})
                    max_points += 1 
                
        summary_table.append({
            "section": checklist_name,
            "points": points,
            "max_points": max_points
            })
        total_points += max_points
    
    summary_form["total"] = total_points
    summary_form["table"] = summary_table
    
    return form_with_ans, summary_form

def add_image_to_word(cell, image_filename):
    #img_name, img_extension = os.path.splitext(image_filename)
    image_iobyte = download_image(image_filename, "singhealth")
    
    im = io.BytesIO()
    image = Image.open(image_iobyte)
    #image.rotate(270).save(im, format=image.format)
    
    add_image = cell.paragraphs[0]
    run = add_image.add_run()
    run.add_picture(image_iobyte, width = 4000000)
    
def add_rectification_form(document, form_with_ans):
    hex_color = form_with_ans["color"]
    rect_form = form_with_ans["rect_form"]
    
    if rect_form != {}:
        header4 = document.add_paragraph()
        runner = header4.add_run(form_with_ans["type"])
        runner.underline = True
        for section in rect_form.keys():
            rect = document.add_table(rows=1, cols=2)
            rect.style = 'TableGrid'
            hdr_cells = rect.rows[0].cells
            hdr_cells[0].text = section
            set_cell_background(hdr_cells[0], hex_color)
            hdr_cells[1].text = 'Description'
            set_cell_background(hdr_cells[1], hex_color)

            for question in rect_form[section].keys():
                row_cells = rect.add_row().cells
                row_cells[0].merge(row_cells[1])
                row_cells[0].text = question
                set_cell_background(row_cells[0], "FDFD96")
                
                info = rect_form[section][question]
                for item in info:
                    row_cells = rect.add_row().cells
                    row_cells[0].text = item
                    
                    if (item != "Non-compliance Images") and (item != "Rectification Images"):
                        row_cells[1].text = str(info[item])
                    
                    elif (item == "Non-compliance Images"):
                        if info[item] is not None:
                            for i in info[item]:
                                add_image_to_word(row_cells[1], i)
                                
                        else:
                            row_cells[1].text = "-"
                            
                    else:
                        if info[item] is not None:
                            for i in info[item]:
                                add_image_to_word(row_cells[1], i)
                                
        set_column_width(rect.columns[0], Inches(1.57))
        set_column_width(rect.columns[1], Inches(4.6))

def add_summary_to_docs(document, form_with_ans):
    header4 = document.add_paragraph()
    runner = header4.add_run(form_with_ans["type"])
    runner.underline = True

    hex_color = form_with_ans["color"]
    summary_form = form_with_ans["summary_form"]
    total = summary_form["total"]
    
    summary = document.add_table(rows=1, cols=2)
    hdr_cells = summary.rows[0].cells
    hdr_cells[0].text = "Section"
    set_cell_background(hdr_cells[0], hex_color)
    hdr_cells[1].text = 'Points'
    set_cell_background(hdr_cells[1], hex_color)

    weighted_total_score = 0
    weighted_max_score = 0
    
    if form_with_ans["type"] == "F&B":
        for section in summary_form["table"]:
            row_cells = summary.add_row().cells
            sectionName = section["section"]
            row_cells[0].text = sectionName
            
            if sectionName == "Professionalism and Staff Hygiene":
                weighted_max_score = 10
            elif sectionName == "Housekeeping and General Cleanliness":
                weighted_max_score = 20
            elif sectionName == "Food Hygiene":
                weighted_max_score = 35
            elif sectionName == "Healthier Choice in line with HPB's Healthy Eating's Initiative":
                weighted_max_score = 15
            elif sectionName == "Workplace Safety and Health":
                weighted_max_score = 20
            else:
                header5 = document.add_paragraph("***Error in generating word docs")
                return
                
            score = section["points"] / section["max_points"] 
            weighted_score = score * weighted_max_score
            row_cells[1].text = "{:.2f}\t /{:.2f}%".format(weighted_score, weighted_max_score)
            weighted_total_score += weighted_score
            
    
    elif form_with_ans["type"] == "Non-F&B":
        for section in summary_form["table"]:
            row_cells = summary.add_row().cells
            sectionName = section["section"]
            row_cells[0].text = sectionName
            
            if sectionName == "Professionalism and Staff Hygiene":
                weighted_max_score = 20
            elif sectionName == "Housekeeping and General Cleanliness":
                weighted_max_score = 40
            elif sectionName == "Workplace Safety and Health":
                weighted_max_score = 40
            else:
                header5 = document.add_paragraph("***Error in generating word docs")
                return
            
            score = section["points"] / section["max_points"]
            weighted_score = score * weighted_max_score
            row_cells[1].text = "{:.2f}\t /{:.2f}%".format(weighted_score, weighted_max_score)
            weighted_total_score += weighted_score
            
    
    elif form_with_ans["type"] == "Covid-19":
        for section in summary_form["table"]:
            row_cells = summary.add_row().cells
            row_cells[0].text = section["section"]

            weighted_score = section["points"] / total * 100
            weighted_max_score = section["max_points"] / total * 100
            row_cells[1].text = "{:.2f}\t /{:.2f}%".format(weighted_score, weighted_max_score)
            weighted_total_score += weighted_score

    summary_cells = summary.add_row().cells
    summary_cells[0].text = "Total:"
    summary_cells[1].text = "{:.2f}\t /100.0%".format(weighted_total_score)
        
    set_column_width(summary.columns[0], Inches(4.7))
    set_column_width(summary.columns[1], Inches(1.7))
    make_rows_bold(summary.rows[-1])

def add_form_to_docs(document, form_with_ans):
    document.add_page_break()
    header3 = document.add_paragraph()
    runner = header3.add_run(form_with_ans["type"])
    runner.underline = True

    hex_color = form_with_ans["color"]

    for checklist_name in form_with_ans["form_with_ans"].keys():
        checklist = document.add_table(rows=1, cols=2)
        checklist.style = 'TableGrid'

        hdr_cells = checklist.rows[0].cells
        hdr_cells[0].text = checklist_name
        set_cell_background(hdr_cells[0], hex_color)
        hdr_cells[1].text = 'Point(s) \nAwarded'
        set_cell_background(hdr_cells[1], hex_color)

        for question in form_with_ans["form_with_ans"][checklist_name]:
            row_cells = checklist.add_row().cells
            row_cells[0].text = question["question"]
            row_cells[1].text = question["answer"]

        set_column_width(checklist.columns[0], Inches(5.4))
        set_column_width(checklist.columns[1], Inches(0.76))
        table_center_align(checklist.columns[1])
        document.add_paragraph("* Rectified")

def from_audit_to_word(document, data):
    audit_dict = data["audit_info"]
    inst_dict = data["inst_info"]
    staff_dict = data["staff_info"]
    tenant_dict = data["tenant_info"]
    all_checklist = audit_dict["form_with_ans"]

    # unpack data
    score = round(audit_dict["score"] * 100, 2)
    date = audit_dict["date"]

    # def fields
    staff_data = ["name", "email"]
    inst_data = ["name", "_id"]
    tenant_data = ["name", "email", "stallName"]

    # Define font style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Default fields
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = title.add_run('SINGHEALTH RETAIL - AUDIT REPORT')
    runner.bold = True
    runner.underline = True

    # Audit information
    header1 = document.add_paragraph()
    runner = header1.add_run("Audit Information")
    runner.bold = True
    runner.underline = True

    # def value
    info = document.add_table(rows=13, cols=2)
    def_info = ["Staff", "Name:", "Email:", "Institution name:", "Institution acronym:", "",
                "Tenant", "Name :", "Email:", "Stall name:", "",
                "Audit date:", "Total score:"]
    insert_list_by_col(info, 0, 0, def_info)
    set_column_width(info.columns[0], Inches(1.6))
    make_rows_underline(info.rows[0], info.rows[6])

    # insert data
    insert_dict_by_col(info, 1, 1, staff_data, staff_dict)
    insert_dict_by_col(info, 3, 1, inst_data, inst_dict)
    insert_dict_by_col(info, 7, 1, tenant_data, tenant_dict)
    info.cell(11, 1).text = date.strftime('%Y/%m/%d %I:%M %p')
    info.cell(12, 1).text = str(score)+"%"
    
     # Add checklists
    for form in all_checklist.values():
        add_form_to_docs(document, form)
    
    
    # Add rectification
    if score < 100:
        document.add_page_break()
        header4 = document.add_paragraph()
        runner = header4.add_run("Rectification Information")
        runner.underline = True
        for form in all_checklist.values():
            add_rectification_form(document, form)

    #  summary
    document.add_page_break()
    header4 = document.add_paragraph()
    runner = header4.add_run("Summary")
    runner.underline = True
    for form in all_checklist.values():
        add_summary_to_docs(document, form)
        document.add_paragraph("\n")

#For emails
def send_audit_email_word(app, to_email, subject, message,
                          word_name, data):

    # for mailing
    app.config.update(BTS_APP_CONFIG_EMAIL)
    mail = Mail(app)

    try:
        msg = Message(sender=BTS_EMAIL,
                      recipients=[to_email],
                      subject=subject,
                      body=message)

        # make temp file
        fd, path = tempfile.mkstemp(suffix=".docx")

        # make word
        document = Document()
        from_audit_to_word(document, data)

        document.save(path)

        with open(path, 'rb') as f:
            file_data = f.read()

        # add as email attachment
        msg.attach(word_name+".docx",
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document", file_data)

        mail.send(msg)

        # close temp file after email sent to delete temp file
        os.close(fd)

    except:
        return False

    return True

def send_audit_email_excel(app, to_email, subject, message,
                           excel_name, page_name, data):

    # for mailing
    app.config.update(BTS_APP_CONFIG_EMAIL)
    mail = Mail(app)

    try:
        msg = Message(sender=BTS_EMAIL,
                      recipients=[to_email],
                      subject=subject,
                      body=message)

        # make temp file
        fd, path = tempfile.mkstemp(suffix=".xlsx")

        # make excel
        workbook = xlsxwriter.Workbook(path)
        from_audit_to_excel(workbook, page_name, data)

        with open(path, 'rb') as f:
            file_data = f.read()

        # add as email attachment
        msg.attach(excel_name+".xlsx",
                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", file_data)

        mail.send(msg)

        # close temp file after email sent to delete temp file
        os.close(fd)
    except:
        return False

    return True

def send_email_notif(app, to_email, subject, message):
    # for mailing
    # for mailing
    app.config.update(BTS_APP_CONFIG_EMAIL)
    mail = Mail(app)

    try:
        msg = Message(
            sender=BTS_EMAIL,
            recipients=[to_email],
            subject=subject,
            body=message
        )
        mail.send(msg)
    except:
        print("Mail failed to send")